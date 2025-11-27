import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
import io
import os
import re
import random

# --- Helper Functions and Constants ---

def generate_vibrant_rgb_colors(count=200):
    """Generates a list of highly saturated, distinct RGB colors (BROADER SPECTRUM for diversity)."""
    colors = set()
    while len(colors) < count:
        h = random.random()
        s = 0.9 # Saturation cao
        v = 0.8 # Value/Brightness cao (sử dụng dải màu rộng hơn)
        
        if s == 0.0: r = g = b = v
        else:
            i = int(h * 6.0); f = h * 6.0 - i; p = v * (1.0 - s); q = v * (1.0 - s * f); t = v * (1.0 - s * (1.0 - f))
            if i % 6 == 0: r, g, b = v, t, p
            elif i % 6 == 1: r, g, b = q, v, p
            elif i % 6 == 2: r, g, b = p, v, t
            elif i % 6 == 3: r, g, b = p, q, v
            elif i % 6 == 4: r, g, b = t, p, v
            else: r, g, b = v, p, q
        
        r, g, b = int(r * 255), int(g * 255), int(b * 255)
        # Không giới hạn màu tối/sáng để tối đa hóa sự đa dạng
        colors.add((r, g, b))
    
    return list(colors)

FONT_COLORS_RGB_200 = generate_vibrant_rgb_colors(200) # Sử dụng 200 màu
speaker_color_map = {}
used_colors = []

def get_speaker_color(speaker_name):
    """Assigns unique, diverse color (Font RGB) to a speaker."""
    global used_colors
    global speaker_color_map
    
    if speaker_name not in speaker_color_map:
        if used_colors:
            color_object = used_colors.pop()
        else:
            r, g, b = random.choice(FONT_COLORS_RGB_200) # Sử dụng pool 200
            color_object = RGBColor(r, g, b)
            
        speaker_color_map[speaker_name] = color_object
        
        # ĐÃ LOẠI BỎ HIGHLIGHT HOÀN TOÀN
        
    return speaker_color_map[speaker_name]

# List of common phrases mistakenly identified as speakers (for filtering)
NON_SPEAKER_PHRASES = {
    "AND REMEMBER", "OFFICIAL DISTANCE", "GOOD NEWS FOR THEIR TEAMMATES", 
    "LL BE HONEST", "FIRST AND FOREMOST", "I SAID", "THE ONLY THING LEFT TO SETTLE", 
    "QUESTION IS", "FINALISTS", "CONTESTANTS", "TEAM PURPLE", "TEAM GREEN", 
    "TEAM PINK", "DUDE PERFECT", "TITLE VO", "WHISPERS", "SRT CONVERSION", 
    "WILL RED THRIVE OR WILL RED BE DEAD", "BUT REMEMBER", "THE RESULTS ARE IN", 
    "WE CHALLENGED", "I THINK", "IN THEIR DEFENSE", "THE PEAK OF HIS LIFE WAS DOING THE SPACETHING",
    "THE ROCKETS ARE BIGGER", "THE DISTANCE SHOULD BE FURTHER", "GET CRAFTY", "THAT WAS SO SICK",
    "OUT OF 100 CONTESTANTS", "THE FIRST ROUND IS BRUTAL", "YOU KNOW WHICH END GOES",
    "THE GAME IS ON", "THAT'S A GOOD THROW", "HE'S GOING FOR IT", "WE GOT THIS",
    "LAUNCH", "OH NO", "OH", "AH", "YEP", "WAIT", "YEAH", "WOO", "OKAY", "YES"
}

# Regexes remain the same
SPEAKER_REGEX_DELIMITER = re.compile(r"([A-Z][a-z\s&]+):\s*", re.IGNORECASE)
TIMECODE_REGEX = re.compile(r"^\d{2}:\d{2}:\d{2},\d{3}\s+-->\s+\d{2}:\d{2}:\d{2},\d{3}$")
HTML_CONTENT_REGEX = re.compile(r"((?:</?[ibu]>)+)(.*?)(?:</?[ibu]>)+", re.IGNORECASE | re.DOTALL)

def set_all_text_formatting(doc, start_index=0):
    """Applies Times New Roman 12pt and general paragraph formatting."""
    for i, paragraph in enumerate(doc.paragraphs):
        if i < start_index:
            continue
            
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12) 
        
        # Line Spacing 1.5 Lines, 0pt Before/After
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0) 

def apply_html_formatting_to_run(paragraph, current_text):
    """Adds text to a paragraph, applying Bold/Italic if enclosed in <i>/<b>/<u> tags."""
    if not current_text.strip():
        return
        
    matches = list(HTML_CONTENT_REGEX.finditer(current_text))
    last_end = 0
    
    for match in matches:
        tag_text = match.group(2) 
        start, end = match.span()

        if start > last_end:
            paragraph.add_run(current_text[last_end:start])
        
        run_html = paragraph.add_run(tag_text)
        run_html.font.bold = True
        run_html.font.italic = True
        
        last_end = end

    if last_end < len(current_text):
        paragraph.add_run(current_text[last_end:])

# Logic for handling Tabs and Indentation (Hanging Indent structure)
def format_and_split_dialogue(document, text):
    """
    Splits a raw text line (which might contain multi-speakers) into separate dialogue 
    paragraphs and applies the required Tab/Hanging Indent formatting.
    """
    
    parts = SPEAKER_REGEX_DELIMITER.split(text)
    TAB_STOP_POSITION = Inches(1.0) # Dialogue start position
    
    # ---------------------------------------------
    # CASE 1: NO SPEAKER FOUND (Continuation Line)
    # ---------------------------------------------
    if len(parts) == 1:
        new_paragraph = document.add_paragraph()
        
        # Áp dụng cấu trúc Hanging Indent
        new_paragraph.paragraph_format.left_indent = TAB_STOP_POSITION
        new_paragraph.paragraph_format.first_line_indent = Inches(-1.0) 
        new_paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.LEFT)
        
        new_paragraph.add_run('\t') # Luôn chỉ dùng 1 Tab cho nội dung tiếp tục
        
        # Spacing
        new_paragraph.paragraph_format.space_after = Pt(0) 
        new_paragraph.paragraph_format.space_before = Pt(0)
        
        apply_html_formatting_to_run(new_paragraph, text)
        return
    
    # ---------------------------------------------
    # CASE 2: ONE OR MORE SPEAKERS FOUND
    # ---------------------------------------------

    # Iterate through all identified speakers
    speaker_matches = list(SPEAKER_REGEX_DELIMITER.finditer(text))
    last_processed_index = 0
    
    for i, match in enumerate(speaker_matches):
        speaker_full = match.group(0) 
        speaker_name = match.group(1).strip()
        start, end = match.span()
        
        # 1. Process Leading Content (content before the current speaker)
        leading_content = text[last_processed_index:start].strip()
        if leading_content:
            continuation_paragraph = document.add_paragraph()
            continuation_paragraph.paragraph_format.left_indent = TAB_STOP_POSITION
            continuation_paragraph.paragraph_format.first_line_indent = Inches(-1.0)
            continuation_paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.LEFT)
            
            continuation_paragraph.add_run('\t') 
            continuation_paragraph.paragraph_format.space_after = Pt(0)
            continuation_paragraph.paragraph_format.space_before = Pt(0)
            apply_html_formatting_to_run(continuation_paragraph, leading_content)

        # Check for Non-Speaker Phrase
        if speaker_name.upper() in NON_SPEAKER_PHRASES:
            content_block = text[start:] 
            
            continuation_paragraph = document.add_paragraph()
            continuation_paragraph.paragraph_format.left_indent = TAB_STOP_POSITION
            continuation_paragraph.paragraph_format.first_line_indent = Inches(-1.0)
            continuation_paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.LEFT)
            
            continuation_paragraph.add_run('\t') 
            apply_html_formatting_to_run(continuation_paragraph, content_block)
            continuation_paragraph.paragraph_format.space_after = Pt(0)
            continuation_paragraph.paragraph_format.space_before = Pt(0)
            return # Exit function as the rest of the line is handled

        # Process Valid Speaker
        
        # Determine the content belonging to this speaker
        if i + 1 < len(speaker_matches):
            next_match_start = speaker_matches[i+1].start()
        else:
            next_match_start = len(text)
            
        content = text[end:next_match_start].strip()

        new_paragraph = document.add_paragraph()
        
        # Áp dụng cấu trúc Hanging Indent cho tất cả các dòng đối thoại
        new_paragraph.paragraph_format.left_indent = TAB_STOP_POSITION
        new_paragraph.paragraph_format.first_line_indent = Inches(-1.0)
        
        # Set Tab Stop at 1.0 inch
        new_paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.LEFT)
        
        # 1. Run cho tên người nói (Bold và Color)
        font_color_object = get_speaker_color(speaker_name) 
        run_speaker = new_paragraph.add_run(speaker_full)
        run_speaker.font.bold = True
        run_speaker.font.color.rgb = font_color_object 
        
        # 2. Xử lý Tab Linh hoạt (1 Tab hoặc 2 Tabs)
        if len(speaker_full) > 10:
             new_paragraph.add_run('\t\t') 
        else:
             new_paragraph.add_run('\t') 

        # 3. Thêm nội dung (NẰM TRÊN CÙNG DÒNG VỚI TÊN NGƯỜI NÓI)
        if content:
            apply_html_formatting_to_run(new_paragraph, content)

        # Spacing
        new_paragraph.paragraph_format.space_after = Pt(0)
        new_paragraph.paragraph_format.space_before = Pt(0)
        
        last_processed_index = next_match_start # Cập nhật vị trí xử lý cuối cùng
    
    # Process remaining content after the last speaker
    remaining_content = text[last_processed_index:].strip()
    if remaining_content:
        continuation_paragraph = document.add_paragraph()
        continuation_paragraph.paragraph_format.left_indent = TAB_STOP_POSITION
        continuation_paragraph.paragraph_format.first_line_indent = Inches(-1.0)
        continuation_paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.LEFT)
        continuation_paragraph.add_run('\t') 
        continuation_paragraph.paragraph_format.space_after = Pt(0)
        continuation_paragraph.paragraph_format.space_before = Pt(0)
        apply_html_formatting_to_run(continuation_paragraph, remaining_content)
        
    return 

# --- Main Processing Function ---

def process_docx(uploaded_file, file_name_without_ext):
    
    global speaker_color_map
    global used_colors
    global highlight_map 
    
    # Reset maps and shuffle color pool for unique assignment per file run
    speaker_color_map = {}
    highlight_map = {} 
    used_colors_rgb = [RGBColor(r, g, b) for r, g, b in FONT_COLORS_RGB_200] # Sử dụng pool 200
    random.shuffle(used_colors_rgb)
    used_colors = used_colors_rgb 
    
    original_document = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in original_document.paragraphs]
    
    document = Document()
    
    # 1. Title Paragraph (Size 20)
    title_text_raw = file_name_without_ext.upper()
    title_text = title_text_raw.replace("CONVERTED_", "").replace("FORMATTED_", "").replace("_EDIT", "").replace(" (GỐC)", "").strip()
    
    title_paragraph = document.add_paragraph(title_text)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0) 
    
    title_run = title_paragraph.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20) 
    title_run.bold = True
    
    # 2. Collect and Add Speaker List (Size 12)
    unique_speakers_ordered = []
    seen_speakers = set()
    
    for paragraph in raw_paragraphs:
        text = paragraph.text
        if text.lower().startswith("srt conversion") or text.lower().startswith("converted_"):
             continue 
             
        for match in SPEAKER_REGEX_DELIMITER.finditer(text):
            speaker_name = match.group(1).strip()
            # Filter non-speaker names
            if speaker_name.upper() not in NON_SPEAKER_PHRASES and speaker_name not in seen_speakers:
                seen_speakers.add(speaker_name)
                unique_speakers_ordered.append(speaker_name)
            
    if unique_speakers_ordered:
        speaker_list_text = "VAI: " + ", ".join(unique_speakers_ordered) 
        speaker_list_paragraph = document.add_paragraph(speaker_list_text)
        
        for run in speaker_list_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12) 
            run.font.bold = False
        
        speaker_list_paragraph.paragraph_format.space_after = Pt(6) 
        speaker_list_paragraph.paragraph_format.space_before = Pt(0)
    
    # 3. Add 2 blank lines
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    
    start_index_for_general_format = len(document.paragraphs)

    # --- B. Process Dialogue Content ---
    
    for paragraph in raw_paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # FIX: Loại bỏ dòng văn bản thừa có cùng tên với Tiêu đề
        if text.upper() == title_text.upper():
            continue
        
        if text.lower().startswith("srt conversion") or text.lower().startswith("converted_"):
            continue 
            
        # Remove Line Numbers
        if re.fullmatch(r"^\s*\d+\s*$", text):
            continue 
            
        # Timecode
        if TIMECODE_REGEX.match(text):
            new_paragraph = document.add_paragraph(text)
            for run in new_paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman' 
                run.font.size = Pt(12) 
            new_paragraph.paragraph_format.space_after = Pt(6) 
            new_paragraph.paragraph_format.space_before = Pt(0) 
            
        # Dialogue Content 
        else:
            format_and_split_dialogue(document, text)
            
    # C. Apply Global Formatting (1.5 Lines)
    for paragraph in document.paragraphs[start_index_for_general_format:]:
        # Apply 1.5 Lines Spacing
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
    # Save the file
    modified_file = io.BytesIO()
    document.save(modified_file)
    modified_file.seek(0)
    
    return modified_file

# --- File Naming and Streamlit UI (English Version) ---

def clean_file_name_for_output(original_filename):
    """Strips unwanted prefixes/suffixes and adds '_edit'."""
    name_without_ext = os.path.splitext(original_filename)[0]
    
    cleaned_name = name_without_ext.replace("CONVERTED_", "").replace("FORMATTED_", "").strip()
    cleaned_name = re.sub(r'\s*\(.*\)$', '', cleaned_name).strip() 
    
    if cleaned_name.lower().endswith("_edit"):
         cleaned_name = cleaned_name[:-5].strip()

    return f"{cleaned_name}_edit.docx"

st.set_page_config(page_title="Automatic Word Script Editor", layout="wide")

st.markdown("## 📄 Automatic Subtitle Script (.docx) Converter")
st.markdown("A Python/Streamlit tool to automatically format subtitle scripts based on specific requirements.")
st.markdown("---")

uploaded_file = st.file_uploader(
    "1. Upload your Word file (.docx)",
    type=['docx'],
    help="Only Microsoft Word .docx format is accepted."
)

if uploaded_file is not None:
    original_filename = uploaded_file.name
    file_name_without_ext = os.path.splitext(original_filename)[0] 
    
    st.info(f"File received: **{original_filename}**.")
    
    if st.button("2. RUN AUTOMATIC FORMATTING"):
        with st.spinner('Processing and formatting Word file...'):
            try:
                modified_file_io = process_docx(uploaded_file, file_name_without_ext)
                
                new_filename = clean_file_name_for_output(original_filename)

                st.success("✅ Formatting complete! You can download the file.")
                
                st.download_button(
                    label="3. Download Formatted Word File",
                    data=modified_file_io,
                    file_name=new_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.markdown("---")
                st.balloons()

            except Exception as e:
                st.error(f"An error occurred during processing: {e}")
                st.warning("Please check the formatting of your input file (e.g., ensure it is a valid DOCX and not corrupted).")
